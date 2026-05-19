import re
import hashlib
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional
from datetime import datetime


@dataclass
class RawDocument:
    content: str
    source: str
    file_type: str
    page: Optional[int] = None
    total_pages: Optional[int] = None
    doc_id: str = field(default="")
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.md5(
                f"{self.source}:{self.page}:{self.content[:64]}".encode()
            ).hexdigest()[:12]


class DocumentLoader:
    def __init__(self, supported_extensions: Optional[set] = None):
        self.supported_extensions = supported_extensions or {".pdf", ".docx", ".txt", ".md"}

    def load(self, path: str | Path) -> List[RawDocument]:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.supported_extensions}")

        loaders = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_txt,
            ".md": self._load_txt,
        }

        documents = loaders[ext](path)
        for doc in documents:
            doc.metadata["loaded_at"] = datetime.utcnow().isoformat()
            doc.metadata["file_size_bytes"] = path.stat().st_size

        return documents

    def load_directory(self, directory: str | Path, recursive: bool = True) -> List[RawDocument]:
        directory = Path(directory)
        if not directory.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory}")

        pattern = "**/*" if recursive else "*"
        all_docs: List[RawDocument] = []

        for file_path in sorted(directory.glob(pattern)):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    docs = self.load(file_path)
                    all_docs.extend(docs)
                    print(f"  Loaded {len(docs)} page(s) from {file_path.name}")
                except Exception as e:
                    print(f"  Warning: Failed to load {file_path.name}: {e}")

        return all_docs

    def _load_pdf(self, path: Path) -> List[RawDocument]:
        try:
            import fitz
        except ImportError:
            raise ImportError("Install pymupdf: pip install pymupdf")

        documents = []
        with fitz.open(str(path)) as pdf:
            total_pages = len(pdf)
            for page_num, page in enumerate(pdf, start=1):
                text = page.get_text("text")
                text = self._clean_text(text)
                if len(text.strip()) < 10:
                    continue
                documents.append(
                    RawDocument(
                        content=text,
                        source=str(path),
                        file_type="pdf",
                        page=page_num,
                        total_pages=total_pages,
                        metadata={
                            "filename": path.name,
                            "page": page_num,
                            "total_pages": total_pages,
                        },
                    )
                )
        return documents

    def _load_docx(self, path: Path) -> List[RawDocument]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = Document(str(path))
        paragraphs = []
        current_section = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_section:
                    paragraphs.append(" ".join(current_section))
                    current_section = []
            else:
                current_section.append(text)

        if current_section:
            paragraphs.append(" ".join(current_section))

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        full_text = "\n\n".join(p for p in paragraphs if p.strip())
        full_text = self._clean_text(full_text)

        if not full_text.strip():
            return []

        return [
            RawDocument(
                content=full_text,
                source=str(path),
                file_type="docx",
                page=1,
                total_pages=1,
                metadata={"filename": path.name},
            )
        ]

    def _load_txt(self, path: Path) -> List[RawDocument]:
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        content = None

        for enc in encodings:
            try:
                content = path.read_text(encoding=enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue

        if content is None:
            raise ValueError(f"Could not decode file {path.name} with any known encoding")

        content = self._clean_text(content)
        if not content.strip():
            return []

        file_type = "markdown" if path.suffix.lower() == ".md" else "txt"
        return [
            RawDocument(
                content=content,
                source=str(path),
                file_type=file_type,
                page=1,
                total_pages=1,
                metadata={"filename": path.name},
            )
        ]

    def _clean_text(self, text: str) -> str:
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[^\S\n]+\n", "\n", text)
        return text.strip()
